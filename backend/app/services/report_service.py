"""DOCX incident report generator."""
import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict
from uuid import uuid4

import aiofiles
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from app.core.config import settings


class ReportService:
    """Generates professional DOCX incident reports."""

    async def generate_docx(self, event, narrative: Dict[str, Any]) -> str:
        """Generate a DOCX report and return its file path."""
        doc = Document()

        # ── Cover ────────────────────────────────────────────────
        self._set_doc_style(doc)
        self._add_cover(doc, event, narrative)
        self._add_executive_summary(doc, narrative)
        self._add_event_details(doc, event)
        self._add_behavior_analysis(doc, event)
        self._add_threat_assessment(doc, event, narrative)
        self._add_recommended_actions(doc, narrative)
        self._add_footer(doc)

        # ── Save ─────────────────────────────────────────────────
        report_dir = Path(settings.REPORTS_DIR)
        report_dir.mkdir(parents=True, exist_ok=True)
        filename = f"incident_{event.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = report_dir / filename

        # Save in executor
        import asyncio
        loop = asyncio.get_event_loop()
        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        async with aiofiles.open(filepath, "wb") as f:
            await f.write(content)

        return str(filepath)

    def _set_doc_style(self, doc: Document):
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def _add_cover(self, doc: Document, event, narrative: Dict):
        doc.add_heading('SECURITY INCIDENT REPORT', 0)
        doc.add_paragraph(f"Classification: {narrative.get('classification', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
        doc.add_paragraph(f"Event ID: {event.id}")
        doc.add_paragraph(f"Severity: {event.severity.upper()}")
        doc.add_paragraph(f"Threat Score: {event.threat_score:.1%}")
        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, narrative: Dict):
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(narrative.get('summary', 'No summary available.'))

    def _add_event_details(self, doc: Document, event):
        doc.add_heading('Event Details', level=1)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        rows_data = [
            ('Event Type', event.event_type.replace('_', ' ').title()),
            ('Timestamp', str(event.timestamp)),
            ('Camera Zone', event.zone_name or 'Unknown'),
            ('Track ID', str(event.track_id or 'N/A')),
            ('Frame Number', str(event.frame_number or 'N/A')),
            ('Confidence', f"{event.confidence:.1%}"),
            ('Threat Score', f"{event.threat_score:.1%}"),
        ]
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

    def _add_behavior_analysis(self, doc: Document, event):
        doc.add_heading('Behavioral Analysis', level=1)
        if event.behavior_flags:
            for flag in event.behavior_flags:
                doc.add_paragraph(f"• {flag.replace('_', ' ').title()}", style='List Bullet')
        else:
            doc.add_paragraph('No specific behavioral flags detected.')

    def _add_threat_assessment(self, doc: Document, event, narrative: Dict):
        doc.add_heading('Threat Assessment', level=1)
        doc.add_paragraph(narrative.get('severity_justification', ''))
        doc.add_paragraph(f"Confidence Notes: {narrative.get('confidence_notes', '')}")

    def _add_recommended_actions(self, doc: Document, narrative: Dict):
        doc.add_heading('Recommended Actions', level=1)
        for i, action in enumerate(narrative.get('recommended_actions', []), 1):
            doc.add_paragraph(f"{i}. {action}")

    def _add_footer(self, doc: Document):
        doc.add_page_break()
        doc.add_paragraph(
            'This report was generated automatically by the AI Surveillance Intelligence Platform. '
            'All findings should be reviewed by qualified security personnel.',
            style='Caption'
        )
